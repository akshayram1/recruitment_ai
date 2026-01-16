"""Document parsing service for PDF and DOCX files"""
import io
from typing import Optional
import fitz  # PyMuPDF
from docx import Document


class DocumentService:
    """Service for parsing PDF and DOCX documents"""
    
    def __init__(self):
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    async def parse_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        text_parts = []
        
        try:
            # Open PDF from bytes
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            doc.close()
            
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
        
        return "\n\n".join(text_parts)
    
    async def parse_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        text_parts = []
        
        try:
            # Open DOCX from bytes
            doc = Document(io.BytesIO(file_content))
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
        
        return "\n\n".join(text_parts)
    
    async def parse_document(self, file_content: bytes, file_type: str) -> str:
        """Parse document based on file type"""
        file_type = file_type.lower()
        
        if file_type in ["pdf", "application/pdf"]:
            return await self.parse_pdf(file_content)
        elif file_type in ["docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            return await self.parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def chunk_text(self, text: str) -> list[str]:
        """Split text into overlapping chunks for embedding"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for period, newline, or other sentence boundaries
                for sep in [". ", "\n\n", "\n", ". ", "! ", "? "]:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > self.chunk_size // 2:
                        end = start + last_sep + len(sep)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
        
        return chunks
